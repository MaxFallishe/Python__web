import openpyxl
from docx import Document
from docx.shared import Pt
import os
import comtypes.client
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders


def main():

    path = 'xxxxxxxxxxxxxxxxxxxxxxxx'
    wb_obj = openpyxl.load_workbook(path)
    sheet_obj = wb_obj.active

    up_text = sheet_obj.cell(row=2, column=23)
    down_text = sheet_obj.cell(row=2, column=24)
    gmail_login = sheet_obj.cell(row=4, column=23)
    gmail_password = sheet_obj.cell(row=4, column=24)

    for i in range(2, sheet_obj.max_row + 1):
        try:
            email_to_send = sheet_obj.cell(row=i, column=3)
            full_name = sheet_obj.cell(row=i, column=4)
            full_name = full_name.value
            full_name = str(full_name)
            full_name = full_name.split(' ')
            if len(full_name) == 2:
                create_docx_letter("Main_letter_sample.docx", full_name[0] + ' ' + full_name[1])
            else:
                create_docx_letter("Main_letter_sample.docx", full_name[1] + ' ' + full_name[2])
            create_pdf_letter("Edited_template.docx")
            email_sender(email_to_send.value, up_text.value, down_text.value, gmail_login.value, gmail_password.value)
            print('Letter №' + str(i-1) + ' is send')
        except Exception as e:
            print(e)
            print("Error №1. Ошибка при отправке письма №" + str(i-1) +
                  "Проверьте корректность заполнения данных для строчки:" + str(i) + "в таблице " + path)


def create_docx_letter(file_template, full_name):
    document = Document(file_template)
    for paragraph in document.paragraphs:
        if 'Уважаемый(ая)_____________________________!' in paragraph.text:
            paragraph.text = ' '

            style = document.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)

            paragraph.style = document.styles['Normal']

            runner = paragraph.add_run('Здравствуйте, ' + str(full_name) + '!')
            runner.bold = True
            # print(paragraph.text)
            break
    document.save('Edited_template.docx')


def create_pdf_letter(docx_template):
    wd_format_pdf = 17

    in_file = os.path.abspath(docx_template)
    out_file = os.path.abspath('Offer_By_State_Symbols_RK.pdf')  

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wd_format_pdf)
    doc.Close()
    word.Quit()


def email_sender(email_to_send, up_text, down_text, login, password):
    email_user = login  
    email_password = password         
    email_send = email_to_send  

    subject = up_text

    msg = MIMEMultipart()
    msg['From'] = email_user
    msg['To'] = email_send
    msg['Subject'] = subject

    body = down_text

    msg.attach(MIMEText(body, 'plain'))
    filename = 'Offer_By_State_Symbols_RK.pdf '  
    attachment = open(filename, 'rb')

    part = MIMEBase('application', 'octet-stream')
    part.set_payload(attachment.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', "attachment; filename= " + filename)

    msg.attach(part)

    letter = msg.as_string()

    smtpserver = smtplib.SMTP_SSL('smtp.mail.ru', 465)  

    smtpserver.ehlo()

    # smtpserver.starttls()    #only for gmail
    smtpserver.ehlo()
    smtpserver.login(email_user, email_password)

    letter = letter.encode('utf-8')
    smtpserver.sendmail(email_user, email_send, letter)


if __name__ == '__main__':
    main()
